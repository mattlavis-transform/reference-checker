import json
import os
import classes.globals as g
from classes.preference_checker import PreferenceChecker
from docx import Document
from docx.shared import Cm
from docx.enum.section import WD_ORIENT
from docx.oxml.shared import OxmlElement, qn
from docx.table import Table
from classes.cached_table import CachedTable

omit_origins = []
include_origins = []
# omit_origins = ["CH", "FO", "IS"]
# include_origins = ["SG"]

# Step 0 - Get the list of all possible origins
with open(g.app.preferences_config) as jsonFile:
    origins = json.load(jsonFile)

# Step 1 - Get the scope, based on the omit / include variables
if len(omit_origins) > 0 and len(include_origins) == 0:
    for origin in omit_origins:
        origins.pop(origin)

elif len(include_origins) > 0:
    for origin in list(origins):
        if origin not in include_origins:
            origins.pop(origin)

# Step 2 - Write the data to the JSON files
for origin in origins:
    preference_checker = PreferenceChecker(origin, origins[origin])

# Step 3 - Write the data extract to the Word report
filename_template = "resources/template/differences_template.docx"
filename_docx = "resources/differences.docx"
document = Document(filename_template)
document.add_heading('FTA issues', 0)
folder = os.path.join(os.getcwd(), "resources", "preferences", "dest")
all_errors = []
for origin in origins:
    filename = os.path.join(folder, origin + ".json")
    with open(filename) as jsonFile:
        errors = json.load(jsonFile)
        all_errors += errors
        jsonFile.close()

row_count = 0
for item in all_errors:
    if item["ignore"] == False:
        row_count += 1
table = CachedTable.transform(document.add_table(rows=row_count + 1, cols=8))
table.style = "List Table 3"

# 25.89 cm is the total table width
widths = (Cm(2.5), Cm(3), Cm(3.5), Cm(3.5),
          Cm(3.5), Cm(3.5), Cm(3.5), Cm(2.89))
for row in table.rows:
    for idx, width in enumerate(widths):
        row.cells[idx].width = width

hdr_cells = table.rows[0].cells
headers = ["Geo ID", "Country", "Commodity code", "Specific issues",
           "Problem", "Scope", "FTA duty", "Database duty"]
for i in range(0, len(headers)):
    hdr_cells[i].text = headers[i]

g.app.set_repeat_table_header(table.rows[0])

row_count = 1
for item in all_errors:
    if item["ignore"] == False:
        hdr_cells = table.rows[row_count].cells
        hdr_cells[0].text = g.app.process_null(item["origin"])
        hdr_cells[1].text = g.app.process_null(item["country"])
        hdr_cells[2].text = g.app.process_null(item["commodity_code"])
        hdr_cells[3].text = g.app.process_null(item["comm_code_differences"])
        hdr_cells[4].text = g.app.process_null(item["reason"])
        hdr_cells[5].text = g.app.process_null(item["scope"])
        hdr_cells[6].text = g.app.process_null(item["fta_duty_value"])
        hdr_cells[7].text = g.app.process_null(item["db_duty_value"])
        row_count += 1
        print(str(row_count))

document.save(filename_docx)
