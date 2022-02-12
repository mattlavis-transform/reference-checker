import json
import os
from docx import Document
from docx.table import Table

import classes.globals as g
from classes.quota import Quota
from classes.classification import Classification
from classes.database import Database


class ClassificationChecker(object):

    def __init__(self, do_from, do_to):
        self.quotas = []
        
        self.do_from = do_from
        self.do_to = do_to
        
        print("Checking classification")
        self.start_log()
        self.get_classification_from_word_document()

    def start_log(self):
        f = open(g.app.quotas_log, "w+")
        f.close()

    def get_classification_from_word_document(self):
        document = Document(g.app.classification_source)
        chapters_done = 0
        for table in document.tables:
            if len(table.rows[0].cells) == 2:
                if len(table.rows[0].cells[0].text) == 10:
                    chapters_done += 1
                    if chapters_done >= self.do_from:
                        self.classifications = []
                        chapter_id = table.rows[0].cells[0].text[0:2]
                        print("Checking chapter", chapter_id)
                        print("- Reading from Word document")
                        for row in table.rows:
                            if len(row.cells) == 2:
                                goods_nomenclature_item_id = row.cells[0].text
                                if len(row.cells[1].tables) > 0:
                                    tbl = row.cells[1].tables[0]
                                    description = tbl.rows[0].cells[0].text
                                else:
                                    description = row.cells[1].text
                                classification = Classification(goods_nomenclature_item_id, description)
                                self.classifications.append(classification)

                        print("- Reading from database")
                        self.get_chapter_from_db(chapter_id)
                        print("- Comparing data")
                        self.convert_chapter_to_dict(chapter_id)

            if chapters_done > self.do_to - 1:
                break

    def get_chapter_from_db(self, chapter_id):
        d = Database()
        sql = """
        select goods_nomenclature_item_id, producline_suffix, description, number_indents, leaf
        from utils.goods_nomenclature_export_new('""" + chapter_id + """%', '2022-01-01')
        order by goods_nomenclature_item_id, producline_suffix
        """

        g.app.quota_duties_dict = {}
        self.classifications_db = []
        rows = d.run_query(sql)
        for row in rows:
            goods_nomenclature_item_id = row[0]
            productline_suffix = row[1]
            description = row[2]
            number_indents = row[3]
            if number_indents == 0:
                if goods_nomenclature_item_id[-8:] == "00000000":
                    number_indents -= 1
            leaf = row[4]
            if number_indents >= 0:
                classification = Classification(goods_nomenclature_item_id, description, number_indents)
                self.classifications_db.append(classification)

    def convert_chapter_to_dict(self, chapter_id):
        # What's been read from the document
        self.classifications_dict = {}
        self.uniques = {}
        for classification in self.classifications:
            if classification.goods_nomenclature_item_id in self.uniques:
                self.uniques[classification.goods_nomenclature_item_id] += 1
                key = classification.goods_nomenclature_item_id + "_" + \
                    str(self.uniques[classification.goods_nomenclature_item_id])
            else:
                self.uniques[classification.goods_nomenclature_item_id] = 0
                key = classification.goods_nomenclature_item_id + "_0"

            item = {
                "goods_nomenclature_item_id": classification.goods_nomenclature_item_id,
                "number_indents": str(classification.number_indents),
                "description": classification.description
            }
            self.classifications_dict[key] = item

        # What's been read from the database
        self.classifications_dict_db = {}
        self.uniques = {}
        for classification in self.classifications_db:
            if classification.goods_nomenclature_item_id in self.uniques:
                self.uniques[classification.goods_nomenclature_item_id] += 1
                key = classification.goods_nomenclature_item_id + "_" + \
                    str(self.uniques[classification.goods_nomenclature_item_id])
            else:
                self.uniques[classification.goods_nomenclature_item_id] = 0
                key = classification.goods_nomenclature_item_id + "_0"

            item = {
                "goods_nomenclature_item_id": classification.goods_nomenclature_item_id,
                "number_indents": str(classification.number_indents),
                "description": classification.description
            }
            self.classifications_dict_db[key] = item

        # Compare the two
        differences = []
        if self.classifications_dict == self.classifications_dict_db:
            same = True
        else:
            for key, value_doc in self.classifications_dict.items():
                a = 1
                try:
                    value_db = self.classifications_dict_db[key]
                except:
                    value_db = {}
                if value_doc != value_db:
                    item = {
                        "document": value_doc,
                        "db": value_db
                    }
                    differences.append(item)
                    
            # Check for items that are in the DB but not in the doc
            for key, value_db in self.classifications_dict_db.items():
                a = 1
                try:
                    value_doc = self.classifications_dict[key]
                except:
                    value_doc = {}
                    item = {
                        "document": value_doc,
                        "db": value_db,
                        "missing": True
                    }
                    differences.append(item)

        self.write_differences(chapter_id, differences)

    def write_differences(self, chapter_id, differences):
        filename = "differences_{0}.json".format(chapter_id)
        filename = os.path.join(g.app.differences_folder, filename)
        with open(filename, 'w') as f:
            json.dump(differences, f, indent=4)
